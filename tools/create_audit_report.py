from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'NOVA-Commerce-Milestone-1-Audit-Report.docx'

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
LIGHT_GRAY = 'F2F4F7'
INK = '1F2937'
GREEN = '166534'
AMBER = '92400E'
RED = '991B1B'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def font(run, size=11, bold=False, color=INK):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(cell, text, bold=False, color=INK, size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    font(r, size, bold, color)
    set_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    r = p.add_run(text)
    font(r, 16 if level == 1 else 13, True, BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead:
        r = p.add_run(bold_lead)
        font(r, 11, True)
    r = p.add_run(text)
    font(r)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for c, header in zip(table.rows[0].cells, headers):
        set_cell_shading(c, LIGHT_GRAY)
        add_text(c, header, bold=True, color=DARK_BLUE, size=9)
    for row in rows:
        cells = table.add_row().cells
        for c, item in zip(cells, row):
            text, color = item if isinstance(item, tuple) else (item, INK)
            add_text(c, text, color=color, size=9)
    for row in table.rows:
        for c in row.cells:
            set_cell_margins(c)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.paragraph_format.space_after = Pt(0)
r = header.add_run('NOVA COMMERCE  |  MILESTONE 1 DELIVERY AUDIT')
font(r, 8, True, DARK_BLUE)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('Internal delivery assessment  |  20 July 2026')
font(r, 8, False, '6B7280')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run('NOVA Commerce')
font(r, 26, True, DARK_BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Milestone 1 - Foundation & Architecture Setup Audit Report')
font(r, 15, False, BLUE)

meta = doc.add_table(rows=3, cols=2)
meta.style = 'Table Grid'
set_table_geometry(meta, [2700, 6660])
for label, value, row in [
    ('Audit scope', 'Milestone 1 requirements supplied for NOVA Commerce', 0),
    ('Assessment basis', 'Repository structure and implementation files present in the workspace', 1),
    ('Assessment date', '20 July 2026', 2),
]:
    set_cell_shading(meta.cell(row, 0), LIGHT_GRAY)
    add_text(meta.cell(row, 0), label, bold=True, color=DARK_BLUE)
    add_text(meta.cell(row, 1), value)
doc.add_paragraph()

add_heading(doc, 'Executive assessment')
add_body(doc, 'Milestone 1 is substantially implemented as a foundation. The repository contains a professional pnpm monorepo, separate Next.js and NestJS applications, database and queue foundations, configuration validation, API conventions, and local infrastructure definitions.')
add_body(doc, 'The achieved scope is ready to begin domain development for catalog, authentication, cart, checkout, orders, inventory, administration, and notifications. These business features are intentionally not delivered in this milestone.')
add_body(doc, 'Important verification limitation: ', 'Dependencies were not installed and full build/typecheck was not executed because the package download approval was declined. Therefore, implementation evidence is present, but runtime verification remains open.')

add_heading(doc, 'Status legend')
legend_rows = [
    [('Achieved', GREEN), 'Requirement is implemented with clear repository evidence.'],
    [('Partial', AMBER), 'Foundation exists, but scope or runtime verification is incomplete.'],
    [('Not verified', RED), 'Cannot be confirmed until dependencies install and checks run.'],
]
add_table(doc, ['Status', 'Meaning'], legend_rows, [1800, 7560])

add_heading(doc, 'Requirements traceability')
backend_rows = [
    ['NestJS architecture', ('Achieved', GREEN), 'apps/api has Nest entry point, root module, typed configuration, and build scripts.'],
    ['Modular structure', ('Achieved', GREEN), 'Database, Health, and Queue modules establish domain-oriented boundaries.'],
    ['Database connection', ('Partial', AMBER), 'Prisma service connects/disconnects at lifecycle boundaries; live connection not tested.'],
    ['Prisma setup', ('Achieved', GREEN), 'Prisma schema defines User, Role, and UserStatus foundation.'],
    ['Migration system', ('Achieved', GREEN), 'Initial SQL migration and prisma migrate scripts are present.'],
    ['Config module', ('Achieved', GREEN), 'ConfigModule validates DATABASE_URL, REDIS_URL, API prefix, port, and web origin.'],
    ['Error handling', ('Achieved', GREEN), 'Global exception filter returns predictable error payloads.'],
    ['API response pattern', ('Achieved', GREEN), 'Global interceptor provides success/data/meta envelope.'],
    ['Redis setup', ('Partial', AMBER), 'Redis Compose service and Redis URL-driven BullMQ connection are configured; not run.'],
    ['Queue foundation', ('Achieved', GREEN), 'Notifications queue and processor skeleton are registered with BullMQ.'],
]
add_table(doc, ['Backend requirement', 'Status', 'Evidence / assessment'], backend_rows, [2600, 1300, 5460])

add_heading(doc, 'Frontend traceability')
frontend_rows = [
    ['Next.js architecture / App Router', ('Achieved', GREEN), 'apps/web uses src/app layout and page routes with Next configuration.'],
    ['UI system', ('Partial', AMBER), 'Tailwind v4, utility function, and base Button component exist; full Shadcn component set is not yet established.'],
    ['API client', ('Achieved', GREEN), 'Typed API client handles base URL, JSON envelope, and credentials.'],
    ['State management', ('Achieved', GREEN), 'Zustand cart-store foundation is implemented.'],
    ['Theme setup', ('Achieved', GREEN), 'next-themes provider and light/dark CSS tokens are configured.'],
]
add_table(doc, ['Frontend requirement', 'Status', 'Evidence / assessment'], frontend_rows, [2600, 1300, 5460])

add_heading(doc, 'Infrastructure and developer readiness')
infra_rows = [
    ['Docker foundation', ('Partial', AMBER), 'PostgreSQL and Redis Compose services plus application Dockerfiles are present. API/web are not yet wired into Compose.'],
    ['Environment management', ('Achieved', GREEN), 'Root and app-specific .env.example files document local configuration.'],
    ['Code standards', ('Achieved', GREEN), 'Strict TypeScript settings, ESLint/Prettier configuration, workspace scripts, and Turbo task graph are present.'],
    ['Developer outcome', ('Partial', AMBER), 'Structure is ready for feature modules; install/build/migration verification remains a release gate.'],
]
add_table(doc, ['Area', 'Status', 'Evidence / assessment'], infra_rows, [2600, 1300, 5460])

add_heading(doc, 'Open items before feature work')
open_items = [
    ['1', 'Run pnpm install, pnpm typecheck, pnpm lint, and pnpm build; resolve any generated/runtime issues.'],
    ['2', 'Start Docker services and run Prisma generate/migrate against PostgreSQL; confirm API health endpoint.'],
    ['3', 'Add full Shadcn UI initialization and common primitives before building storefront screens.'],
    ['4', 'Wire API and web services into Compose or add a production deployment pipeline.'],
    ['5', 'Begin Milestone 2 with Auth/RBAC and Catalog domain models, tests, and API contracts.'],
]
add_table(doc, ['Priority', 'Recommended action'], open_items, [1100, 8260])

add_heading(doc, 'Conclusion')
add_body(doc, 'The requested Milestone 1 architecture is materially achieved. The remaining work is primarily validation and a few platform-completeness tasks, not a redesign of the foundation. The recommended gate is to complete the open items above before declaring the foundation production-verified.')

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = 'NOVA Commerce Milestone 1 Audit Report'
doc.core_properties.subject = 'Foundation and architecture delivery assessment'
doc.core_properties.author = 'NOVA Commerce'
doc.save(OUT)
print(OUT)
