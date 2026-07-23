from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NOVA-Commerce-Milestone-1-Verification-Readiness-Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
LIGHT_FILL = "F8FAFC"
GREEN_FILL = "EAF7EA"
AMBER_FILL = "FFF4CC"
RED_FILL = "FCE8E6"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", size=10.5, bold=False, color=None, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_fill(hdr[i], GRAY_FILL)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "NOVA Commerce | Milestone 1 Verification"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, color=RGBColor(90, 90, 90))

    footer = section.footer.paragraphs[0]
    footer.text = "Foundation readiness report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_font(run, size=9, color=RGBColor(90, 90, 90))

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("NOVA Commerce")
    set_font(r, size=24, bold=True, color=RGBColor(0, 0, 0))
    subtitle = add_para(
        doc,
        "Milestone 1 Verification and Production Readiness Report",
        size=14,
        color=RGBColor(70, 70, 70),
        after=10,
    )
    subtitle.paragraph_format.space_after = Pt(12)

    meta = [
        ("Date", date.today().isoformat()),
        ("Scope", "Verification, stabilization, and production readiness assessment"),
        ("Restriction", "No Milestone 2 business features were added"),
        ("Overall status", "Approved with Minor Fixes"),
        ("Final recommendation", "Pending only environment runtime verification after Docker Desktop is available"),
    ]
    add_table(doc, ["Field", "Value"], meta, widths=[2200, 7160])

    add_heading(doc, "1. Overall Status")
    add_para(
        doc,
        "Milestone 1 Status: Approved with Minor Fixes. The repository now passes lint, typecheck, and production build validation. One backend runtime bug and one lint foundation gap were found and fixed. Docker-based runtime verification remains environment-blocked because Docker Engine was not reachable from the machine during the final check.",
    )

    add_heading(doc, "2. Completed Areas")
    completed = [
        ("Architecture", "Monorepo layout is valid with apps/api and apps/web separated through pnpm workspace filters."),
        ("Backend foundation", "NestJS module structure, config validation, global validation pipe, exception filter, response interceptor, health module, Prisma, Redis and BullMQ foundation are present."),
        ("Database", "Prisma schema and migration structure exist. Prisma generate and earlier migration flow were verified; final deploy check was blocked by Docker Engine availability."),
        ("Frontend foundation", "Next.js App Router, TypeScript, Tailwind, reusable UI component base, API client, Zustand store, and theme provider are present."),
        ("Infrastructure", "Docker Compose defines Postgres and Redis services, env examples exist, and local scripts are available for dev/build/lint/typecheck."),
        ("Code quality", "Strict TypeScript checks, ESLint v9 flat config for API, Next ESLint config, and Prettier config are in place."),
    ]
    add_table(doc, ["Area", "Verification Result"], completed, widths=[2200, 7160])

    add_heading(doc, "3. Issues Found and Actions Taken")
    issues = [
        (
            "Backend lint configuration",
            "Medium",
            "API lint failed because ESLint v9 requires a flat eslint.config.* file; only legacy .eslintrc files existed.",
            "Added apps/api/eslint.config.mjs and installed @typescript-eslint/parser plus @typescript-eslint/eslint-plugin as dev dependencies.",
        ),
        (
            "Backend queue runtime startup",
            "High",
            "Production API start crashed because NotificationsProcessor imported QUEUES from queue.module, creating a circular dependency at runtime.",
            "Moved queue names into apps/api/src/queue/queue.constants.ts and updated module/processor imports.",
        ),
        (
            "Frontend lint warning",
            "Low",
            "ESLint warned about anonymous default export in postcss.config.mjs.",
            "Changed PostCSS config to export a named constant.",
        ),
        (
            "Docker runtime verification",
            "Medium",
            "Docker Engine was not reachable at npipe dockerDesktopLinuxEngine, so Postgres/Redis startup and final migrate deploy could not be completed in this run.",
            "Keep Docker Desktop running and rerun docker compose up -d postgres redis, then pnpm --filter @nova/api prisma:deploy.",
        ),
        (
            "Seed structure",
            "Low",
            "No seed script is currently defined. This is not blocking for Milestone 1 because there is no domain seed data yet, but the structure should be added before feature-heavy milestones.",
            "Add a foundation seed entry when auth/catalog baseline data is introduced.",
        ),
    ]
    add_table(doc, ["Issue", "Severity", "Impact", "Fix"], issues, widths=[1900, 1200, 3100, 3160])

    add_heading(doc, "4. Verification Evidence")
    evidence = [
        ("lint", "Passed", "pnpm.cmd lint completed for @nova/api and @nova/web."),
        ("typecheck", "Passed", "pnpm.cmd typecheck completed for API and Web."),
        ("build", "Passed", "pnpm.cmd build completed: Nest build passed and Next.js production build compiled successfully."),
        ("frontend production build", "Passed", "Next.js generated static routes / and /_not-found successfully."),
        ("backend compile", "Passed", "Nest build completed after queue circular dependency fix."),
        ("Prisma generate", "Passed earlier", "Prisma Client generation completed successfully."),
        ("migration deploy", "Environment-blocked", "Failed because Docker/Postgres was unavailable at final verification time."),
        ("application startup", "Partially verified", "Backend runtime crash was fixed; full health endpoint verification depends on Docker Postgres/Redis availability."),
    ]
    add_table(doc, ["Check", "Status", "Evidence"], evidence, widths=[2200, 1800, 5360])

    add_heading(doc, "5. Readiness Assessment")
    readiness = [
        ("Repository & architecture", "Ready", "No major folder restructuring required. Separation between API and Web is maintained."),
        ("Backend foundation", "Ready with runtime dependency note", "Architecture can support Auth, Products, Orders, Inventory, and Payments without redesign."),
        ("Database", "Ready with Docker dependency", "Schema and migrations are coherent; final DB initialization requires Docker Engine running."),
        ("Redis & queues", "Ready after fix", "BullMQ queue registration is stable after moving queue constants out of the module import cycle."),
        ("Frontend foundation", "Ready", "App Router, UI base, state, theme, and API client layers are suitable for future storefront/auth/admin work."),
        ("UI system", "Ready with expected growth", "Base component structure exists; more Shadcn components can be added as features require them."),
        ("Infrastructure", "Pending local Docker check", "Compose config exists, but runtime check could not complete because Docker Engine was unavailable."),
    ]
    add_table(doc, ["Area", "Status", "Assessment"], readiness, widths=[2200, 1900, 5260])

    add_heading(doc, "6. Final Recommendation")
    add_para(
        doc,
        "Recommendation: Approved with Minor Fixes. From a code quality and build perspective, Milestone 1 is stable and ready for feature development. Before officially starting Milestone 2, run the final local runtime gate with Docker Desktop active: docker compose up -d postgres redis, pnpm --filter @nova/api prisma:deploy, pnpm --filter @nova/api start, and verify GET /api/v1/health.",
    )
    add_para(
        doc,
        "No business logic from future milestones was added during this verification phase. Changes were limited to configuration, lint readiness, and a real backend runtime stability bug.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
