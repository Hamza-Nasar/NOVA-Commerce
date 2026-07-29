from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/NOVA-Commerce-Milestone-2-QA-Implementation-Report.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 99, 110)
GREEN = RGBColor(31, 122, 74)
AMBER = RGBColor(122, 90, 0)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "D9DEE7"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    color = BLUE if level < 3 else RGBColor(31, 77, 120)
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_status_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="C8D4E3")
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Milestone 2 Status: APPROVED WITH MINOR QA NOTES")
    set_run_font(r, size=12, color=GREEN, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Authentication, authorization, and user management foundation is verified against the required scope. Ready for next milestone planning after optional polish notes.")
    set_run_font(r2, size=10.5, color=NAVY)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, head in enumerate(headers):
        shade_cell(hdr[idx], LIGHT_GRAY)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(head)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5)
            if value in ("PASS", "APPROVED", "FIXED"):
                set_run_font(r, size=9.5, color=GREEN, bold=True)
            elif value in ("MINOR", "NOTE"):
                set_run_font(r, size=9.5, color=AMBER, bold=True)
    set_table_borders(table)
    if widths:
        set_table_width(table, widths)
    return table


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    for i, size in [(1, 16), (2, 13), (3, 12)]:
        style = styles[f"Heading {i}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if i < 3 else RGBColor(31, 77, 120)
        style.paragraph_format.space_before = Pt(16 if i == 1 else 12 if i == 2 else 8)
        style.paragraph_format.space_after = Pt(8 if i == 1 else 6 if i == 2 else 4)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("NOVA Commerce | Milestone 2 QA")
    set_run_font(hr, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Prepared for implementation sign-off")
    set_run_font(fr, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("NOVA Commerce")
    set_run_font(r, size=23, color=NAVY, bold=True)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sr = sub.add_run("Milestone 2 QA & Implementation Report")
    set_run_font(sr, size=14, color=MUTED, bold=True)
    for label, value in (
        ("Project", "Production Grade Full-Stack E-commerce Platform"),
        ("Milestone", "Milestone 2 - Authentication, Authorization & User Management"),
        ("Report date", "July 29, 2026"),
        ("Assessment type", "API, browser/manual QA, build, and production-readiness verification"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5)

    add_status_callout(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "Milestone 2 has been implemented and verified according to the required authentication, authorization, and user management scope. The foundation includes backend auth APIs, JWT/refresh-token handling, RBAC guard structure, user profile and address management, frontend auth screens, protected routes, API client/session state, and Docker-backed local validation.")
    add_body(doc, "No product catalog, cart, checkout, payment, order, inventory, or other future milestone business features were added during this verification phase.")

    add_heading(doc, "2. Required Scope Coverage", 1)
    add_table(
        doc,
        ["Area", "Status", "Evidence"],
        [
            ("Backend authentication", "PASS", "Register, login, logout, refresh rotation, and /auth/me verified through API and browser flows."),
            ("Authorization foundation", "PASS", "JWT strategy, auth guard, roles decorator/guard, and role enum are present for future protected modules."),
            ("User management", "PASS", "Profile read/update, password-change API foundation, and address CRUD/default-address flow are implemented."),
            ("Frontend auth UX", "PASS", "Login, register, forgot/reset password, profile, settings, and address pages render and connect to the API."),
            ("Session/state foundation", "PASS", "API client, token memory store, refresh retry handling, Zustand auth/user/session/ui stores are available."),
            ("Database foundation", "PASS", "Prisma migration deployed; users, refresh tokens, and user addresses schema are consistent."),
            ("Infrastructure readiness", "PASS", "Docker Desktop, Postgres, Redis, API startup, and frontend production build were verified."),
        ],
        [2300, 1300, 5760],
    )

    add_heading(doc, "3. Validation Evidence", 1)
    add_table(
        doc,
        ["Check", "Result", "Notes"],
        [
            ("Docker Desktop", "PASS", "Docker engine was running; Postgres and Redis containers were healthy/reachable."),
            ("Postgres port 5432", "PASS", "TCP connectivity succeeded."),
            ("Redis port 6379", "PASS", "TCP connectivity succeeded."),
            ("Prisma deploy", "PASS", "Two migrations found; no pending migration issues."),
            ("Lint", "PASS", "Root pnpm lint completed for API and web."),
            ("Typecheck", "PASS", "Root pnpm typecheck completed for API and web."),
            ("Build", "PASS", "API compiled; Next.js production build generated expected app routes."),
            ("API smoke flow", "PASS", "Health, register, me, profile update, address create/list, refresh rotation, and logout passed."),
            ("Browser QA flow", "PASS", "Register-to-profile, settings update, address CRUD, logout redirect, and login again passed."),
        ],
        [2300, 1300, 5760],
    )

    add_heading(doc, "4. Browser and Manual Flow QA", 1)
    add_body(doc, "Browser QA was executed against the local frontend on port 3006 and API on port 4000 with Docker-backed Postgres/Redis. Screenshots were saved under docs/milestone-2-browser-qa for audit traceability.")
    add_table(
        doc,
        ["Flow", "Result", "Detail"],
        [
            ("Login page render", "PASS", "Login form and navigation displayed correctly."),
            ("Registration", "PASS", "New customer registered successfully and redirected to protected profile."),
            ("Profile settings", "PASS", "Profile update request returned 200 and UI reflected updated name."),
            ("Address management", "PASS", "Address created, listed as default, and deleted successfully."),
            ("Logout", "PASS", "Logout endpoint returned success and protected profile redirected to login."),
            ("Login after registration", "PASS", "Registered user could log in again successfully."),
        ],
        [2500, 1300, 5560],
    )

    add_heading(doc, "5. Issues Found and Stabilization Fixes", 1)
    add_table(
        doc,
        ["Issue", "Severity", "Impact", "Fix / Recommendation"],
        [
            ("CORS origin mismatch for QA port 3006", "FIXED", "Browser QA requests were blocked when frontend used an alternate port.", "WEB_ORIGIN handling now supports comma-separated origins and main.ts enables CORS using parsed origin list."),
            ("Empty optional profile fields", "FIXED", "Blank optional phone/profileImage could trigger validation errors.", "Frontend settings submit now sends undefined for empty optional fields."),
            ("Missing favicon", "FIXED", "Browser requested /favicon.ico and produced avoidable 404 noise.", "Added app icon asset."),
            ("Password-form browser warning", "MINOR", "Chrome reports an accessibility/autocomplete warning on password-related forms.", "Add hidden/associated username field around password-change/reset forms in a polish pass."),
            ("Next standalone start warning", "MINOR", "next start warns when output: standalone is configured, although QA server still responded successfully.", "Add a dedicated start:standalone script or adjust production smoke startup to run .next/standalone/server.js with static assets copied correctly."),
        ],
        [2200, 1150, 2500, 3510],
    )

    add_heading(doc, "6. Remaining Risks / Notes", 1)
    for item in [
        "Email delivery, OTP verification, products, cart, checkout, orders, inventory, payments, and admin CMS remain outside Milestone 2 and should not be treated as missing for this phase.",
        "RBAC foundations are in place, but deeper role-permission matrices should be finalized when admin/dashboard modules are introduced.",
        "Refresh-token security is foundation-ready; production deployment should still confirm secure cookie settings, HTTPS, domain, and same-site policy per environment.",
        "The browser QA folder includes a historical diagnostic screenshot from an earlier failed attempt; the final PASS evidence screenshots are the numbered successful flow images.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Final Recommendation", 1)
    add_body(doc, "Recommendation: Approved with minor QA notes. Milestone 2 satisfies the required scope and is ready for sign-off. The team can proceed to the next milestone after optional polish on password-form autocomplete and standalone production start scripting.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A - API Smoke Output", 1)
    for line in [
        "HEALTH=ok",
        "REGISTER_USER=smoke+1785347837@nova.test",
        "ME_USER=smoke+1785347837@nova.test",
        "PROFILE_NAME=Smoke Verified",
        "ADDRESS_COUNT=1",
        "ADDRESS_DEFAULT=True",
        "REFRESH_TOKEN_ROTATED=True",
        "LOGOUT=True",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=9.5)

    add_heading(doc, "Appendix B - Browser QA Pass List", 1)
    for line in [
        "PASS API health",
        "PASS Frontend home",
        "PASS Login page renders",
        "PASS Register redirects to protected profile",
        "PASS Profile update succeeds",
        "PASS Address create/list/default succeeds",
        "PASS Address delete succeeds",
        "PASS Protected route redirects when logged out",
        "PASS Login succeeds after registration",
    ]:
        add_bullet(doc, line)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
